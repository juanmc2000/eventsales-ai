"""Generate a human-readable Word document from response_prep_100_results.json.

Usage:
    python tests/scripts/generate_response_prep_100_docx.py

Reads:
    tests/data/response_prep_100_results.json
    tests/data/freeform_group_booking_response_preparation_test_100.json

Writes:
    tests/data/response_prep_100_results_sprint14b.docx
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# ── Colour palette (matches existing sprint11 doc) ────────────────────────────

C_NAVY      = RGBColor(0x0D, 0x2A, 0x52)
C_DARK      = RGBColor(0x1A, 0x1A, 0x1A)
C_GREY      = RGBColor(0x55, 0x55, 0x55)
C_GREEN     = RGBColor(0x1A, 0x7A, 0x3C)
C_RED       = RGBColor(0xC0, 0x39, 0x2B)
C_ORANGE    = RGBColor(0xE6, 0x7E, 0x22)
C_MONO_BG   = RGBColor(0xF5, 0xF5, 0xF5)

# ── Paths ─────────────────────────────────────────────────────────────────────

REPO_ROOT   = Path(__file__).resolve().parents[2]
RESULTS     = REPO_ROOT / "tests/data/response_prep_100_results.json"
FIXTURE     = REPO_ROOT / "tests/data/freeform_group_booking_response_preparation_test_100.json"
OUT_PATH    = REPO_ROOT / "tests/data/response_prep_100_results_sprint14b.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_run_colour(run, colour: RGBColor) -> None:
    run.font.color.rgb = colour


def _para(doc: Document, text: str, style: str = "Normal",
          colour: RGBColor = C_DARK, bold: bool = False, size_pt: int | None = None) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    _set_run_colour(run, colour)
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)
    return p


def _coloured_line(doc: Document, label: str, label_colour: RGBColor,
                   rest: str, rest_colour: RGBColor = C_DARK) -> None:
    p = doc.add_paragraph(style="Normal")
    r1 = p.add_run(label)
    _set_run_colour(r1, label_colour)
    r1.bold = True
    r2 = p.add_run(rest)
    _set_run_colour(r2, rest_colour)


def _draft_block(doc: Document, text: str) -> None:
    """Render draft text in a shaded mono-style paragraph."""
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    _set_run_colour(run, C_DARK)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_summary_table(doc: Document, results: list[dict]) -> None:
    total = len(results)
    llm_calls = sum(1 for r in results if r["generation_path"] == "llm")
    det_calls = sum(1 for r in results if r["generation_path"] == "deterministic")
    compliance_pass = sum(1 for r in results if r["compliance"]["passed"])
    auto_send_ok = sum(1 for r in results if r["auto_send"]["auto_send_allowed"])
    safety_issues = sum(1 for r in results if r["safety_checks"]["has_issues"])
    total_in = sum(r["input_tokens"] for r in results)
    total_out = sum(r["output_tokens"] for r in results)

    goal_counts: dict[str, int] = {}
    for r in results:
        goal_counts[r["response_goal"]] = goal_counts.get(r["response_goal"], 0) + 1

    rows = [
        ("Total records", str(total)),
        ("LLM calls", str(llm_calls)),
        ("Deterministic calls", str(det_calls)),
        ("Compliance pass", f"{compliance_pass}/{total}  ({compliance_pass/total*100:.1f}%)"),
        ("Auto-send eligible", f"{auto_send_ok}/{total}  ({auto_send_ok/total*100:.1f}%)"),
        ("Safety issues", str(safety_issues)),
        ("Total tokens (in / out)", f"{total_in:,} in  /  {total_out:,} out"),
    ]
    for goal, count in sorted(goal_counts.items()):
        rows.append((f"  {goal}", str(count)))

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    _set_run_colour(run, C_DARK)


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    with open(RESULTS) as f:
        data = json.load(f)
    with open(FIXTURE) as f:
        fixture_raw = json.load(f)

    # Build fixture lookup by id
    fixture: dict[str, dict] = {r["id"]: r for r in fixture_raw["records"]}

    results: list[dict] = data["results"]
    run_id: str = data["run_id"]
    generated_at: str = data["generated_at"][:10]
    prompt_version: int = data["prompt_version"]

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph(style="Title")
    tr = title.add_run("Sprint 14B — Response Preparation Evaluation Report")
    _set_run_colour(tr, C_NAVY)

    sub = doc.add_paragraph(style="Normal")
    sr = sub.add_run(
        f"100-record evaluation  ·  sprint14b  ·  prompt V{prompt_version}"
        f"  ·  run {generated_at}  ·  run_id {run_id}"
    )
    _set_run_colour(sr, C_GREY)

    doc.add_paragraph()

    # ── Pipeline Changes ──────────────────────────────────────────────────────
    h = doc.add_paragraph(style="Heading 2")
    hr = h.add_run("Pipeline Changes (S14B)")
    _set_run_colour(hr, C_NAVY)

    changes = [
        "RESP-063: REQUEST_DATE_CONFIRMATION fully deterministic — copy blocks only, no LLM call",
        "RESP-064: Draft prompt V7 — ONE warmth sentence constraint; V6 archived",
        "RESP-064: Explicit forbidden-phrase list in MANDATORY RULES (excellent choice, perfect for, ideal, etc.)",
        "RESP-065: _strip_provisional_sentences() removes LLM-era provisional language from RDTC clarification questions",
        "Runner fix: PENDING_DATE_CONFIRMATION / INSUFFICIENT_INFORMATION correctly mapped to their own contract values",
    ]
    for c in changes:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(c)
        _set_run_colour(r, C_DARK)

    doc.add_paragraph()

    # ── Summary table ─────────────────────────────────────────────────────────
    h2 = doc.add_paragraph(style="Heading 2")
    h2r = h2.add_run("Summary")
    _set_run_colour(h2r, C_NAVY)

    _add_summary_table(doc, results)

    doc.add_paragraph()

    # ── Email Evaluations ─────────────────────────────────────────────────────
    h1 = doc.add_paragraph(style="Heading 1")
    h1r = h1.add_run("Email Evaluations")
    _set_run_colour(h1r, C_NAVY)

    for idx, record in enumerate(results, 1):
        rec_id      = record["record_id"]
        goal        = record["response_goal"]
        auto_ok     = record["auto_send"]["auto_send_allowed"]
        auto_label  = "AUTO-SEND ✓" if auto_ok else "AUTO-SEND ✗"
        auto_colour = C_GREEN if auto_ok else C_RED
        comp_pass   = record["compliance"]["passed"]

        fix = fixture.get(rec_id, {})
        sender      = fix.get("sender", {})
        sender_name = sender.get("name", "—")
        sender_email= sender.get("email", "—")
        sender_dom  = sender.get("domain", "—")
        body_full   = fix.get("body", record.get("body_preview", ""))

        te          = fix.get("target_extraction", {})
        rpt         = te.get("response_preparation_target", {})
        avd         = rpt.get("availability_decision", {})
        av_status   = avd.get("availability_status", record.get("availability_status", "—"))
        dpv         = rpt.get("draft_prompt_variables", {})
        restaurant  = dpv.get("restaurant_name", "—")
        room_raw    = dpv.get("room_lines", "")
        room        = room_raw.replace("Suggested room:", "").strip().rstrip(".") if room_raw else "—"
        date_ctx    = rpt.get("date_context", {})
        # Try assumed_date, then first candidate_date, then checked_candidate_dates
        candidate_dates = date_ctx.get("candidate_dates") or avd.get("checked_candidate_dates") or []
        event_date  = (date_ctx.get("assumed_date")
                       or (candidate_dates[0] if candidate_dates else None)
                       or "—")
        meal_period = te.get("meal_period") or "—"

        # ── Record heading ────────────────────────────────────────────────────
        h2 = doc.add_paragraph(style="Heading 2")
        num_run = h2.add_run(f"#{idx:03d}  {rec_id}")
        _set_run_colour(num_run, C_NAVY)
        num_run.bold = True
        goal_run = h2.add_run(f"  [{goal.replace('_', ' ')}]  ")
        _set_run_colour(goal_run, C_NAVY)
        as_run = h2.add_run(f"[{auto_label}]")
        _set_run_colour(as_run, auto_colour)
        as_run.bold = True

        # ── Inbound Email ─────────────────────────────────────────────────────
        h3 = doc.add_paragraph(style="Heading 3")
        _set_run_colour(h3.add_run("Inbound Email"), C_NAVY)

        _para(doc, f"Subject: {fix.get('subject', record.get('subject', '—'))}")
        _para(doc, f"From: {sender_name}  <{sender_email}>")
        _para(doc, f"Domain: {sender_dom}")
        _para(doc, body_full[:400] + ("…" if len(body_full) > 400 else ""))

        # ── Availability Context ──────────────────────────────────────────────
        h3b = doc.add_paragraph(style="Heading 3")
        _set_run_colour(h3b.add_run("Availability Context"), C_NAVY)

        _para(doc, f"Restaurant: {restaurant}")
        _para(doc, f"Room: {room}")
        _para(doc, f"Date: {event_date}")
        _para(doc, f"Period: {meal_period}")
        status_col = C_GREEN if av_status == "AVAILABLE" else (C_RED if av_status == "UNAVAILABLE" else C_ORANGE)
        p_status = doc.add_paragraph(style="Normal")
        sr1 = p_status.add_run("Status: ")
        _set_run_colour(sr1, C_DARK)
        sr2 = p_status.add_run(av_status)
        _set_run_colour(sr2, status_col)
        sr2.bold = True

        # ── Draft Response ────────────────────────────────────────────────────
        h3c = doc.add_paragraph(style="Heading 3")
        _set_run_colour(h3c.add_run("Draft Response"), C_NAVY)

        gen_path = record["generation_path"]
        tokens_in  = record["input_tokens"]
        tokens_out = record["output_tokens"]
        token_str  = f"{tokens_in} in / {tokens_out} out" if tokens_in else "deterministic (no tokens)"
        _para(doc, f"Generation path: {gen_path}  ·  {token_str}")

        # Clarification questions (RDTC)
        cq_line = record.get("clarification_questions_line", "")
        if cq_line.strip():
            _para(doc, f"Clarification question: {cq_line.strip()}", colour=C_GREY)

        _draft_block(doc, record.get("response", "—"))

        # ── Compliance ────────────────────────────────────────────────────────
        h3d = doc.add_paragraph(style="Heading 3")
        _set_run_colour(h3d.add_run("Evaluation"), C_NAVY)

        violations = record["compliance"].get("violations", [])
        if comp_pass:
            _coloured_line(doc, "✓ COMPLIANCE PASS", C_GREEN, "  All checks passed")
        else:
            for v in violations:
                _coloured_line(doc, "✗ VIOLATION  ", C_RED, v)

        safety_issues = record["safety_checks"].get("issues", [])
        if safety_issues:
            for s in safety_issues:
                _coloured_line(doc, "⚠ SAFETY  ", C_ORANGE, str(s))

        # Auto-send
        blockers = record["auto_send"].get("auto_send_blockers", [])
        if auto_ok:
            _coloured_line(doc, "Auto-send: ", C_DARK, "ELIGIBLE", C_GREEN)
        else:
            blocker_str = "  ·  ".join(blockers) if blockers else record["auto_send"].get("review_required_reason", "—")
            _coloured_line(doc, "Auto-send: ", C_DARK, f"BLOCKED — {blocker_str}", C_RED)

        doc.add_paragraph()

    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
